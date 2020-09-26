from docx import Document
import csv
import os
# from wand.image import Image

# Simple saver
def save_as_txt_or_html(content: str, path='', file_name='default', add=False):
    path = path if not path or path.endswith('/') else f'{path}/'
    if not os.path.exists(path): os.makedirs(path)

    if file_name.endswith('.txt') or file_name.endswith('.html'):
        doc_path_name = f'{path}{file_name}'
    else:
        doc_path_name = f'{path}{file_name}.txt'

    if os.path.exists(doc_path_name) and add:
        with open(doc_path_name, mode='a', encoding='utf-8') as file:
            file.write(content)
    else:
        with open(doc_path_name, mode='w', encoding='utf-8') as file:
            file.write(content)


# Best visual representation
def save_as_doc(content: list, path='', file_name='default', add=False):
    path = path if not path or path.endswith('/') else f'{path}/'
    if not os.path.exists(path): os.makedirs(path)
    doc_path_name = f'{path}{file_name}' if file_name.endswith('.docx') \
        else f'{path}{file_name}.docx'
    if os.path.exists(doc_path_name):
        doc = Document(doc_path_name)
    else:
        doc = Document()
    if not doc.paragraphs or add:
        doc.add_heading(content[0], 0)  # program name
        doc.add_paragraph(content[1])  # program link
        for tutor in content[2]:  # tutors
            # Don`t work because of non-working method
            """image_path = f'{tutor.img_path}{tutor.name}.png'
            with Image(filename=image_path) as img:
                doc.add_picture(img)"""
            doc.add_paragraph(tutor.description())

    doc.save(doc_path_name)

# Can be entered into database
def save_as_csv(content: list, path='', file_name='default', add=False):
    path = path if not path or path.endswith('/') else f'{path}/'
    if not os.path.exists(path): os.makedirs(path)
    doc_path_name = f'{path}{file_name}' if file_name.endswith('.csv') \
        else f'{path}{file_name}.csv'

    if os.path.exists(doc_path_name) and add:
        with open(doc_path_name, mode='a', encoding='utf-8') as file:
            writer = csv.writer(file)
            for tutor in content:
                writer.writerow(tutor.csv_entry())
    else:
        with open(doc_path_name, mode='w', encoding='utf-8') as file:
            writer = csv.writer(file)
            for tutor in content:
                writer.writerow(tutor.csv_entry())

# Image loader
# Convenient to take out into a separate function
def degree_images_loager(degree: dict):
    for tutors in degree.values():
        for tutor in tutors:
            tutor.get_image()
# Not all images was loaded because of bad connection
